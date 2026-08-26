import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def markdown_to_docx(md_path, docx_path, title):
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Malgun Gothic'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(15, 23, 42)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    def flush_table():
        nonlocal in_table, table_lines
        if not table_lines:
            in_table = False
            return
        
        # Parse markdown table
        rows = []
        for line in table_lines:
            if re.match(r'^\s*\|?\s*[-:]+[-| :]*\s*$', line):
                continue # delimiter row
            cols = [c.strip() for c in line.strip().split('|')]
            if cols and cols[0] == '': cols.pop(0)
            if cols and cols[-1] == '': cols.pop()
            if cols:
                rows.append(cols)
        
        if rows:
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            
            for r_idx, row_data in enumerate(rows):
                row = table.rows[r_idx]
                is_header = (r_idx == 0)
                for c_idx in range(num_cols):
                    cell = row.cells[c_idx]
                    text = row_data[c_idx] if c_idx < len(row_data) else ""
                    cell.text = text
                    
                    # Style cell
                    if is_header:
                        set_cell_background(cell, "0F766E")
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.size = Pt(10)
                    else:
                        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                        set_cell_background(cell, bg_color)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(30, 41, 59)
            doc.add_paragraph() # space after table

        table_lines = []
        in_table = False

    for raw_line in lines:
        line = raw_line.rstrip()

        # Check table
        if line.strip().startswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        elif in_table:
            flush_table()

        if not line.strip():
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 118, 110) # Brand Dark Teal
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:].strip())
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(2, 132, 199) # Blue Accent
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:].strip())
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(71, 85, 105)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[2:].strip())
            run.font.italic = True
            run.font.color.rgb = RGBColor(51, 65, 85)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            clean_text = line[2:].strip()
            # Handle bold inside bullet
            parts = re.split(r'(\*\*.*?\*\*)', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            clean_text = re.sub(r'^\d+\.\s', '', line).strip()
            parts = re.split(r'(\*\*.*?\*\*)', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
        elif line.startswith('```'):
            continue # skip code fence tags
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Generated docx: {docx_path}")

if __name__ == "__main__":
    base_dir = r"C:\Users\SSAFY\Desktop\Full Course\app"
    
    # 1. LOCAL_ROUTE_기획서_v5.docx
    md1 = os.path.join(base_dir, "LOCAL_ROUTE_기획서_v5.md")
    docx1 = os.path.join(base_dir, "LOCAL_ROUTE_기획서_v5.docx")
    markdown_to_docx(md1, docx1, "LOCAL ROUTE 서비스 기획서 v5")

    # 2. LOCAL_ROUTE_6인_개발계획서_및_업무지침서.docx
    md2 = os.path.join(base_dir, "LOCAL_ROUTE_6인_개발계획서_및_업무지침서.md")
    docx2 = os.path.join(base_dir, "LOCAL_ROUTE_6인_개발계획서_및_업무지침서.docx")
    markdown_to_docx(md2, docx2, "LOCAL ROUTE 6인 개발계획서 및 업무지침서")
